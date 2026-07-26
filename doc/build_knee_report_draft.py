from pathlib import Path
from copy import deepcopy
import hashlib
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

sys.stdout.reconfigure(encoding="utf-8")

ROOT = Path(r"C:\Projects\Study\School-safety-accident")
REFERENCE = ROOT / "doc" / "_report_work" / "reference.docx"
OUTPUT = ROOT / "doc" / "학교안전사고_무릎중대화_보고서초안.docx"
IMG = ROOT / "output_knee"

EXPECTED_REFERENCE_HASH = "B09980F0C6F6016BFAF69405CA019186F48AE8068DE298B50434742A13715D70"
actual_hash = hashlib.sha256(REFERENCE.read_bytes()).hexdigest().upper()
if actual_hash != EXPECTED_REFERENCE_HASH:
    raise RuntimeError(f"Reference changed: {actual_hash}")

doc = Document(REFERENCE)

# Clear the copied body while retaining the reference section properties,
# styles, numbering, theme, and document settings.
body = doc._element.body
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)

section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin = Inches(1.18)
section.right_margin = Inches(1.18)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.header_distance = Inches(0.49)
section.footer_distance = Inches(0.49)

doc.core_properties.title = "학교안전사고 무릎 중대화 요인 분석 및 예방 방안"
doc.core_properties.subject = "고등학교 무릎 장해급여 발생 사고와 신경근 준비운동"
doc.core_properties.author = ""
doc.core_properties.last_modified_by = ""

FONT = "휴먼명조"
ACCENT = RGBColor(197, 54, 42)
GRAY = RGBColor(90, 90, 90)


def set_run_font(run, size=10.3, bold=False, color=RGBColor(0, 0, 0)):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def set_paragraph(
    paragraph,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line=0.27,
    left=0.0,
    right=0.0,
    before=0,
    after=3,
    line=1.55,
    keep_next=False,
):
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Inches(first_line) if first_line is not None else None
    fmt.left_indent = Inches(left)
    fmt.right_indent = Inches(right)
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_next
    fmt.widow_control = True


def add_text(text, size=10.3, bold=False, color=RGBColor(0, 0, 0), **kwargs):
    p = doc.add_paragraph()
    set_paragraph(p, **kwargs)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    return p


def add_major(number, title):
    p = doc.add_paragraph()
    set_paragraph(
        p,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=None,
        left=0,
        before=5,
        after=7,
        line=1.2,
        keep_next=True,
    )
    r1 = p.add_run(f"{number}.")
    set_run_font(r1, 15.5)
    r2 = p.add_run(f"\t{title}")
    set_run_font(r2, 15.5, bold=True)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(0.55))
    return p


def add_sub(number, title):
    p = doc.add_paragraph()
    set_paragraph(
        p,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=None,
        left=0.18,
        before=4,
        after=3,
        line=1.25,
        keep_next=True,
    )
    r1 = p.add_run(f"{number}) ")
    set_run_font(r1, 11.6)
    r2 = p.add_run(title)
    set_run_font(r2, 11.6, bold=True)
    return p


def add_minor(number, title):
    p = doc.add_paragraph()
    set_paragraph(
        p,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=None,
        left=0.32,
        before=3,
        after=2,
        line=1.2,
        keep_next=True,
    )
    r = p.add_run(f"({number}) {title}")
    set_run_font(r, 10.8, bold=True)
    return p


def add_callout(label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(5.75)
    cell = table.cell(0, 0)
    cell.width = Inches(5.75)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F4F5")
    tc_pr.append(shd)
    tc_mar = OxmlElement("w:tcMar")
    for side, value in (("top", "130"), ("left", "180"), ("bottom", "130"), ("right", "180")):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), value)
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)
    p = cell.paragraphs[0]
    set_paragraph(p, first_line=None, left=0, right=0, before=0, after=0, line=1.35)
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, 10.2, bold=True, color=ACCENT)
    r2 = p.add_run(text)
    set_run_font(r2, 10.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def set_alt_text(inline_shape, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def add_figure(filename, caption, note=None, width=5.05):
    p = doc.add_paragraph()
    set_paragraph(
        p,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=None,
        before=2,
        after=2,
        line=1.0,
        keep_next=True,
    )
    shape = p.add_run().add_picture(str(IMG / filename), width=Inches(width))
    set_alt_text(shape, caption)
    cap = add_text(
        caption,
        size=9.5,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=None,
        before=1,
        after=1,
        line=1.15,
        keep_next=bool(note),
    )
    if note:
        add_text(
            note,
            size=8.0,
            color=GRAY,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line=None,
            before=0,
            after=3,
            line=1.15,
        )
    return p


def set_cell_text(cell, text, size=8.8, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph(p, alignment=align, first_line=None, left=0, right=0, before=0, after=0, line=1.2)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, value in (("top", "95"), ("left", "110"), ("bottom", "95"), ("right", "110")):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), value)
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = int(sum(widths) * 1440)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_page_break():
    p = doc.add_paragraph()
    p.add_run().add_break()
    p._p.xpath(".//w:br")[0].set(qn("w:type"), "page")


# PAGE 1 — problem, scope, and research question
title = doc.add_paragraph()
set_paragraph(
    title,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    first_line=None,
    before=4,
    after=5,
    line=1.15,
    keep_next=True,
)
r = title.add_run("학교안전사고 무릎 중대화 요인 분석 및 예방 방안")
set_run_font(r, size=17.5, bold=True)

subtitle = doc.add_paragraph()
set_paragraph(
    subtitle,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    first_line=None,
    before=0,
    after=12,
    line=1.2,
    keep_next=True,
)
r = subtitle.add_run("— 고등학교 비접촉성 무릎 손상과 신경근 준비운동을 중심으로 —")
set_run_font(r, size=10.8, color=GRAY)

add_callout(
    "핵심 결론",
    "무릎의 문제는 단순 발생빈도가 아니라 장해급여로 전환되는 비율이다. "
    "고등학교 무릎 보상사고 10,836건 중 245건(2.26%)에서 장해급여가 발생했으며, "
    "비접촉 상황이 무릎 중대사고의 58.8%를 차지했다. 따라서 기존의 달리기·스트레칭 중심 준비운동을 "
    "근력·균형·착지·감속·방향전환과 교사 피드백을 포함한 표준 신경근 준비운동으로 보완할 필요가 있다.",
)

add_major(1, "연구 배경 및 문제 정의")
add_sub(1, "분석 배경")
add_text(
    "학교안전사고 예방은 흔한 사고의 발생 건수를 줄이는 것뿐 아니라, 사고가 장기 치료와 장해급여로 "
    "이어지는 경로를 차단하는 데에도 초점을 두어야 한다. 본 분석은 2021~2025년 학교안전공제 보상자료에서 "
    "고등학교 사고를 분리한 뒤, 사고부위별 장해급여 발생률을 비교하여 무릎을 핵심 중대화 부위로 선정하였다. "
    "이후 접촉유형, 당시활동, 장소, 사고형태를 교차하여 예방 가능한 상황을 좁혔다."
)

add_sub(2, "조작적 정의와 분석 범위")
add_text(
    "본 무릎 분석에서 ‘중대화’는 해당 보상사고에서 장해급여가 발생한 경우로 조작적으로 정의하였다. "
    "중대화율은 ‘장해급여 발생 사고 건수 ÷ 해당 집단의 전체 보상사고 건수 × 100’으로 계산하였다. "
    "이는 행정자료상의 지표로서 임상적 손상 중증도 전체나 실제 영구장해 발생률과 동일하지 않다. "
    "분석대상은 고등학교 무릎 보상사고 10,836건이며, 이 중 장해급여 발생 사고는 245건이다."
)

add_sub(3, "연구 질문")
add_text(
    "첫째, 무릎 사고의 장해급여 발생률은 다른 부위와 비교해 얼마나 높은가? 둘째, 무릎 중대사고는 직접 충돌보다 "
    "비접촉 동작에 집중되는가? 셋째, 어떤 활동과 사고상황을 우선 관리해야 하는가? 넷째, 최신 의학근거는 "
    "현재의 일반 준비운동과 무엇이 다르며 학교가 어떤 내용을 표준화해야 하는가?"
)

add_text(
    "자료: 학교안전공제 보상자료(2021~2025) 재분석, 스포츠안전재단 2024 공개 원자료, "
    "청소년 신경근 손상예방 프로그램 관련 무작위시험·메타분석.",
    size=8.4,
    color=GRAY,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line=None,
    before=5,
    after=0,
    line=1.25,
)

add_page_break()

# PAGE 2 — why knee
add_major(2, "무릎은 왜 우선 관리 대상인가")
add_sub(1, "부위별 중대화율 비교")
add_figure(
    "01_부위별_중대화율.png",
    "그림 1. 고등학교 보상사고의 주요 사고부위별 중대화율(장해급여 발생률)",
    "주: 점선은 고등학교 전체 보상사고의 평균 중대화율 0.35%이다. 괄호 안 n은 부위별 전체 보상사고 건수이다.",
    width=5.15,
)
add_text(
    "무릎 사고는 10,836건 중 245건에서 장해급여가 발생하여 중대화율이 2.26%였다. 이는 고등학교 전체 "
    "보상사고 평균 0.35%의 약 6.5배이며, 비교한 주요 사고부위 가운데 가장 높았다. 흉부 1.34%(6/449), "
    "허리 1.12%(12/1,073)도 평균보다 높았지만 무릎은 비율뿐 아니라 장해급여 발생 건수에서도 가장 큰 부담을 보였다."
)

add_sub(2, "빈도와 중대화율을 분리한 해석")
add_text(
    "이 결과는 ‘무릎 사고가 가장 자주 발생한다’는 뜻이 아니라, 보상자료에 포함된 무릎 사고가 장해급여로 "
    "이어지는 비율이 높다는 뜻이다. 정책 우선순위는 비율과 절대 건수를 함께 보아야 한다. 표본이 작은 흉부와 "
    "허리는 비율 변동성이 큰 반면, 무릎은 1만 건이 넘는 표본에서도 높은 신호가 유지되어 우선적인 심층분석이 타당하다."
)

add_callout(
    "해석 원칙",
    "“무릎 사고는 다른 부위보다 정확히 6.5배 위험하다”가 아니라 "
    "“이 고등학교 보상자료에서 무릎의 장해급여 발생률이 전체 평균의 약 6.5배였다”라고 표현한다.",
)

add_sub(3, "예방 가능한 경로를 따로 찾아야 하는 이유")
add_text(
    "무릎 장해급여 사고에는 직접 충돌, 비접촉성 염좌, 골절, 과사용 등 서로 다른 손상기전이 섞여 있을 수 있다. "
    "따라서 무릎이라는 부위만으로 해결책을 정하지 않고, 접촉유형과 사고상황을 분해하여 신경근 준비운동이 직접 "
    "겨냥할 수 있는 비접촉성·간접접촉성 동작의 비중을 확인하였다."
)

add_page_break()

# PAGE 3 — why non-contact
add_major(3, "왜 비접촉 동작을 바꾸어야 하는가")
add_sub(1, "전체사고와 중대사고의 접촉유형 구성")
add_figure(
    "02_무릎_접촉유형_구성비.png",
    "그림 2. 고등학교 무릎 전체사고와 중대사고의 접촉유형 구성비",
    "주: 비접촉은 사람·물체와 직접 충돌하지 않고 착지·감속·방향전환 등의 동작 중 발생한 사고로 분류하였다.",
    width=4.75,
)
add_text(
    "비접촉 사고는 전체 무릎 사고의 44.8%였지만 무릎 중대사고에서는 58.8%로 14.0%p 증가하였다. "
    "반면 대인접촉은 전체의 33.9%에서 중대사고의 18.4%로 감소하였다. 이는 충돌 주의만 강조하는 예방교육으로는 "
    "무릎 중대사고의 다수를 설명하기 어렵고, 학생 자신의 착지·급정지·방향전환 능력을 다루어야 함을 보여준다."
)

add_sub(2, "같은 접촉유형 안에서도 무릎의 중대화율이 높음")
add_figure(
    "03_접촉유형별_중대화율.png",
    "그림 3. 고등학교 보상사고의 접촉유형별 중대화율: 무릎과 전체 부위 비교",
    "주: 비율비는 교란요인을 보정하지 않은 기술통계이며 개인의 위험배수나 인과효과가 아니다.",
    width=4.75,
)
add_text(
    "비접촉 무릎 사고의 중대화율은 2.97%로 무릎 접촉유형 중 가장 높았으며, 같은 비접촉 유형의 전체 부위 "
    "중대화율 0.61%의 약 4.9배였다. 대인접촉(무릎 1.23%, 전체 0.19%)과 물체접촉(2.19%, 0.18%)에서도 "
    "무릎의 비율이 높았다. 그림 2와 3을 종합하면 비접촉 상황은 중대사고 내 비중과 중대화율이 모두 높은 핵심 예방대상이다."
)

add_page_break()

# PAGE 4 — priority sports and scenarios
add_major(4, "어디서 무엇을 할 때 우선 개입해야 하는가")
add_sub(1, "활동별 중대화율과 절대 부담")
add_figure(
    "04_무릎_중대화율_사고당시활동.png",
    "그림 4. 고등학교 무릎 사고의 당시활동별 중대화율",
    "주: 비율과 함께 전체 사고 건수 및 장해급여 발생 건수를 고려해 정책 우선순위를 판단하였다.",
    width=4.85,
)
add_text(
    "중대화율은 태권도·유도·합기도가 4.64%(14/302)로 가장 높았고 축구 3.09%(107/3,465), 농구 "
    "2.81%(39/1,387) 순이었다. 그러나 축구는 장해급여 발생 107건으로 무릎 중대사고 245건의 43.7%를 차지했다. "
    "따라서 ‘최고 비율은 무도 종목, 최대 부담은 축구’로 구분하고, 보편적 적용의 1순위는 축구·농구 등 구기활동으로 정한다."
)

add_sub(2, "대표 사고 시나리오")
add_figure(
    "08_무릎중대사고_형태x활동.png",
    "그림 5. 고등학교 무릎 중대사고의 사고형태×당시활동 조합",
    "자료: 학교안전공제 보상자료 재분석. 수치는 장해급여 발생 사고 건수이다.",
    width=4.85,
)
add_text(
    "축구×넘어짐이 49건으로 가장 많았고, 농구×넘어짐 20건, 축구×사람과의 부딪힘 18건이 뒤를 이었다. "
    "이는 축구 활동 전 착지·감속·방향전환 훈련을 우선 도입하되, 충돌과 시설물 주변 위험관리도 함께 시행해야 함을 뜻한다. "
    "장소별로는 운동장에 무릎 중대사고 130건이 집중되었고 실·내외 체육시설도 중대화율 3.66%로 높아, "
    "표준 준비운동은 운동장·체육시설에서 이루어지는 구기수업과 스포츠클럽부터 적용하는 것이 합리적이다."
)

add_callout(
    "데이터가 가리키는 정책 대상",
    "고등학교 축구·농구 활동 / 운동장·실내외 체육시설 / 넘어짐으로 기록된 비접촉·간접접촉 상황.",
)

add_page_break()

# PAGE 5 — evidence and solution
add_major(5, "최신 의학근거를 반영한 해결방안")
add_sub(1, "현재 준비운동과 실제로 무엇이 다른가")
add_text(
    "스포츠안전재단 2024 공개 원자료에서 13~18세 축구 부상 경험자 40명 중 39명(97.5%)이 부상 전 "
    "준비운동을 했다고 답했다. 표본이 작고 전문체육인 비중이 높다는 한계가 있지만, ‘준비운동을 하지 않아서 다친다’는 "
    "가설은 지지되지 않는다. 문제는 준비운동 실시 여부가 한 발 균형, 점프 착지, 감속, 방향전환, 자세 피드백과 "
    "반복 용량을 포함하는지 확인하지 못한다는 데 있다."
)
add_text(
    "일반 준비운동과 신경근 준비운동은 가벼운 달리기와 동적 관절운동에서는 겹친다. 핵심 차이는 "
    "① 둔근·햄스트링·종아리·몸통 근력, ② 한 발 균형, ③ 점프 착지·급정지·방향전환 기술, "
    "④ 교사의 즉각적 자세 피드백, ⑤ 주 2~3회의 반복이다. 따라서 ‘신경근 활성화’라는 새 체조보다 "
    "반복적인 다요소 운동기술 학습이라는 표현이 정확하다."
)

add_sub(2, "효과크기")
add_text(
    "11~15세 학생 725명의 학교 체육수업 군집무작위시험에서 일반 달리기·스트레칭 대비 신경근 프로그램의 "
    "무릎 염좌 발생률비는 0.36(95% 신뢰구간 0.13~0.98)이었다. 2025년 청소년·젊은 선수의 무작위시험 "
    "24편 메타분석에서는 무릎손상이 28% 감소했다(RR 0.72, 95% 신뢰구간 0.62~0.84). "
    "2025년 여자 팀스포츠 메타분석에서는 전체 무릎손상 22%, ACL 손상 50% 감소가 보고됐다. "
    "ACL의 50~60% 감소는 주로 청소년 여자 구기선수 근거이므로 학교 무릎 전체에 그대로 적용하지 않고, "
    "본 보고서의 대표 기대효과는 넓은 무릎손상 약 22~28% 감소로 제시한다."
)

add_sub(3, "10~15분 표준 신경근 준비운동")
rows = [
    ("2분", "동적 달리기", "전진·후진 달리기, 사이드 스텝"),
    ("2~3분", "근력·몸통", "스쿼트, 런지, 브리지·플랭크"),
    ("2분", "한 발 균형", "한 발 지지+뻗기 또는 공 주고받기"),
    ("2~3분", "착지·감속", "작은 점프 후 정지, 앞·옆 점프 착지"),
    ("2~3분", "방향전환", "셔플 후 정지, 계획형→반응형 컷"),
]
table = doc.add_table(rows=1 + len(rows), cols=3)
set_table_geometry(table, [0.7, 1.25, 3.55])
headers = ("시간", "구성", "예시 동작")
for ci, text in enumerate(headers):
    set_cell_text(table.cell(0, ci), text, size=8.9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.cell(0, ci), "E7E8E9")
for ri, row in enumerate(rows, start=1):
    for ci, text in enumerate(row):
        set_cell_text(
            table.cell(ri, ci),
            text,
            size=8.6,
            align=WD_ALIGN_PARAGRAPH.CENTER if ci < 2 else WD_ALIGN_PARAGRAPH.LEFT,
        )
table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))

add_text(
    "교사 피드백은 “무릎은 두 번째 발가락 방향”, “엉덩이와 무릎을 굽혀 조용히 착지”, "
    "“먼저 속도를 줄이고 방향을 바꾸기”, “몸통이 한쪽으로 무너지지 않게”로 통일한다. "
    "해결책은 일회성 재교육이 아니라 실습형 교사 워크숍, 동작 영상·QR, 수업별 구성요소 체크리스트를 묶어 "
    "기존 준비운동을 표준 프로그램으로 교체하는 방식이어야 한다.",
    before=3,
)

add_sub(4, "한계와 최종 제언")
add_text(
    "장해급여는 임상 중증도 전체와 같지 않고, 내부자료는 세부 진단과 활동 노출시간을 제공하지 않는다. "
    "또한 외부 예방연구는 여자 구기선수 비중이 높아 모든 학생과 모든 무릎 손상에 같은 효과를 보장하지 않는다. "
    "따라서 ‘프로그램이 장해급여 사고를 50~60% 줄인다’고 단정하지 않는다. 그럼에도 내부자료에서 비접촉 상황과 "
    "축구×넘어짐이 명확히 집중되고, 학교 기반 무작위시험에서 무릎 염좌 감소가 확인되었으므로, "
    "고등학교 구기활동의 준비운동을 내용·용량·피드백 중심으로 표준화하는 정책 제안은 충분히 타당하다."
)

add_text(
    "주요 출처: Richmond et al.(2016), Clin J Sport Med, PMID 27367045; "
    "LaBella et al.(2011), Arch Pediatr Adolesc Med, PMID 22065184; "
    "Li & Zhu(2025), The Knee, PMID 40618548; Gu et al.(2025), Ann Med, PMID 41175154; "
    "Exercise & Sports Science Australia Position Statement(2026), DOI 10.1007/s40279-026-02450-3; "
    "Lutz et al.(2024), Br J Sports Med, DOI 10.1136/bjsports-2023-106906.",
    size=7.7,
    color=GRAY,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line=None,
    before=3,
    after=0,
    line=1.15,
)

doc.save(OUTPUT)
print(OUTPUT)
