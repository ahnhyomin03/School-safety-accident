from pathlib import Path
import sys

from docx import Document
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding="utf-8")
SOURCE = Path(r"C:\Users\user\Downloads\학교안전사고_중대화요인분석_보고서.docx")
doc = Document(SOURCE)

print(f"sections={len(doc.sections)} paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} inline_shapes={len(doc.inline_shapes)}")
for i, section in enumerate(doc.sections):
    print(
        f"SECTION {i}: size={section.page_width}x{section.page_height} "
        f"margins={section.top_margin},{section.right_margin},{section.bottom_margin},{section.left_margin} "
        f"header={section.header_distance} footer={section.footer_distance}"
    )
    print(" HEADER:", " | ".join(p.text for p in section.header.paragraphs))
    print(" FOOTER:", " | ".join(p.text for p in section.footer.paragraphs))

for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.replace("\t", " ⇥ ").replace("\n", " ↵ ")
    if text or paragraph._p.xpath(".//w:drawing|.//w:br"):
        style_name = paragraph.style.name if paragraph.style is not None else "(none)"
        print(f"P{i:03d} [{style_name}] {text}")

for ti, table in enumerate(doc.tables):
    table_style = table.style.name if table.style is not None else "(none)"
    print(f"TABLE {ti}: {len(table.rows)}x{len(table.columns)} style={table_style}")
    for ri, row in enumerate(table.rows):
        values = [cell.text.replace("\n", " / ") for cell in row.cells]
        print(f"  R{ri}: " + " || ".join(values))

for si, shape in enumerate(doc.inline_shapes):
    print(f"SHAPE {si}: {shape.width}x{shape.height}")

used = {}
for paragraph in doc.paragraphs:
    style_name = paragraph.style.name if paragraph.style is not None else "(none)"
    used[style_name] = used.get(style_name, 0) + 1
print("STYLES_USED:", used)

for style_name in sorted(name for name in used if name != "(none)"):
    style = doc.styles[style_name]
    font = style.font
    pf = style.paragraph_format
    print(
        f"STYLE {style_name}: font={font.name} size={font.size} bold={font.bold} "
        f"align={pf.alignment} before={pf.space_before} after={pf.space_after} "
        f"line={pf.line_spacing} left={pf.left_indent} first={pf.first_line_indent}"
    )
    rpr = style.element.find(qn("w:rPr"))
    if rpr is not None:
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is not None:
            print(" RFONTS", dict(rfonts.attrib))
